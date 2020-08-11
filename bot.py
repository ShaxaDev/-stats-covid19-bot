import telebot 
import os
from docx import Document
from docx.shared import Inches

# ------- initial code -------

document = Document()


token='985403265:AAFSOSsxtLT2FEC_Hq37mrgmRyxIWydSd4Q'
bot=telebot.TeleBot(token)
@bot.message_handler(commands=['ruy'])
def ruy(ms):

	send=bot.send_message(ms.chat.id,'👨‍💼ismingizni kiriting')
	bot.register_next_step_handler(send,ruy1)
def ruy1(ms):
	suz=ms.text
	document.add_heading('Mijoz @{} malumotlari'.format(ms.from_user.username),level=0)
	document.add_heading('Ism : '+suz,level=1)
	
	s=bot.send_message(ms.chat.id,'🏠manzilingiz:')
	bot.register_next_step_handler(s,ruy2)
def ruy2(ms):
	suz1=ms.text
	document.add_heading('Manzil: '+suz1,level=1)
	t=bot.send_message(ms.chat.id,'☎️Telefon raqamingizni kiriting:')
	bot.register_next_step_handler(t,ruy3)
def ruy3(ms):
	try:
		suz2=ms.text
		document.add_heading('Telefon raqami: '+'{}'.format(suz2),level=1)
		t1=bot.send_message(ms.chat.id,'🤵Rasmingizni yuboring: ')
		bot.register_next_step_handler(t1,ruy4)
	except:
		pass
def ruy4(ms):
	document.add_heading('Foydalanuvchi rasmi: ',level=1)
	w=ms.photo[0].file_id
	file_info = bot.get_file(ms.photo[0].file_id)
	download = bot.download_file(file_info.file_path)
	with open('rasm.jpg', 'wb') as new_file:
		new_file.write(download)
	document.add_picture('rasm.jpg')
	
	document.save('{}.docx'.format(ms.from_user.username))
	q=open('{}.docx'.format(ms.from_user.username),'rb')
	bot.send_document(995951832,q)
	q.close()
	os.remove('{}.docx'.format(ms.from_user.username))
	bot.send_message(ms.chat.id,'Ro`yxatdan o`tish muvafaqiyatli yakunlandi✅')








bot.polling()


